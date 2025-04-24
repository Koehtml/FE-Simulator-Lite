import requests
from docx import Document
import os

#API details
API_URL = " "
API_KEY = " " #Potentially optional

#Prompt for AI
PROMPT = "Enter Prompt"

#Function to call API and get AI response
def get_response(prompt):
    headers = {
        "authorization": f"Bearer {API_KEY}", #Only necessary if API needs key
        "Content-Type": "application/json" #Specifies saving in JSON format
    }
    data = {"prompt": prompt} #Tells AI what to respond to

    response = requests.post(API_URL, json=data, headers=headers) #Sends POST request with prompt

    #If request is successfull (Code 100), extract reponse from JSON file. Otherwise error is indicated for debugging
    if response.status_code == 200:
        return response.json().get("response", "No response received")
    else:
        return f"Error: {response.status_code} - {response.text}"
    
#Function to save response to a Word doc with new paragraph for each new response
#Opens the document if it is found, otherwise creates a new document
def save_to_word(text, filename="ai_response.docx"):
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    response_count = len(doc.paragraphs) + 1 #Counts existing responses to number correctly
    
    #Adds new response in sequence
    doc.add_paragraph(f"Response {response_count} saved to {filename}")

    #Saved document
    doc.save(filename)
    print(f"Response {response_count} saved to {filename}")

#Executes functions
ai_response = get_response(PROMPT) #Calls API with prompt
save_to_word(ai_response) #Saves response